"""
Lampiran 7C ASPI Automation Script
===================================
Script untuk mengotomasi pembuatan dokumen Lampiran 7C ASPI (format Word)
dari sumber data UAT Script (format Excel).

Fungsi utama:
- Membaca file UAT Script.xlsx dari folder input/
- Menyalin seluruh isi kolom Remarks langsung ke kolom Request di Lampiran 7C
- Membuat dokumen Word Lampiran 7C dari scratch dengan format yang sesuai
- Menyimpan hasil ke folder output/
"""

import os
import re
import sys
from pathlib import Path

from openpyxl import load_workbook
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ============================================================
# STYLING TABEL (disesuaikan dengan contoh Lampiran 7C)
# ============================================================

# Font default seluruh dokumen
FONT_NAME = "Calibri"

# Label pada kolom Request/Response yang dibuat BOLD (hanya labelnya, isi normal)
BOLD_LABELS = ["URL Endpoint:", "Header Request:", "Request Body:", "Response Body:"]

# Warna latar header tabel (oranye) - hex tanpa tanda pagar
HEADER_FILL_COLOR = "ED7D31"

# Warna blok (shading) untuk baris yang Tidak dites / Belum dites (N/A) - kuning
NA_FILL_COLOR = "FFFF00"

# Lebar tiap kolom dalam TWIPS (1/20 poin), urut: No, Service, Scenario,
# Expected Result, Request, Response, Result, Notes.
# Lebar kolom PERSIS seperti contoh Lampiran 7C (total 15168 twips ~26.7 cm),
# cocok untuk halaman LANDSCAPE.
COLUMN_WIDTHS_TWIPS = [700, 1418, 1984, 1701, 4829, 1975, 851, 1710]

# Kolom yang isinya di-rata-tengah (selain itu rata kiri).
# Indeks: 0=No, 1=Service, 2=Scenario, 3=Expected Result, 6=Result
CENTER_ALIGNED_COLUMNS = {0, 1, 2, 3, 6}


# ============================================================
# KONFIGURASI & KONSTANTA
# ============================================================

# Path konfigurasi
INPUT_DIR = Path("input")
OUTPUT_DIR = Path("output")
UAT_SCRIPT_FILENAME = "UAT Script.xlsx"
OUTPUT_FILENAME = "Lampiran 7C - Hasil UAT.docx"

# Sheet name di file UAT Script
UAT_SHEET_NAME = "UAT Script"

# Info header dokumen Lampiran 7.C
NAMA_PENYEDIA_LAYANAN = "Bank Sahabat Sampoerna"
NAMA_PENGGUNA_LAYANAN = ""
TANGGAL_PENGUJIAN = ""

# Kolom index di UAT Script (0-based, dengan kolom A kosong)
# Kolom A (index 0): KOSONG
# Kolom B (index 1): Kategori Tes
# Kolom C (index 2): Nama Modul
# Kolom D (index 3): Nomor Skenario
# Kolom E (index 4): Nomor Kasus Tes (contoh: 1.1, 2.3, 7.15)
# Kolom F (index 5): Langkah Tes
# Kolom G (index 6): Hasil yang diharapkan
# Kolom H (index 7): Hasil Aktual
# Kolom I (index 8): Remarks (berisi log URL, Headers, Request Body, Response)
# Kolom J (index 9): Tanggal Pelaksanaan
# Kolom K (index 10): Jenis Script
# Kolom L (index 11): Pelaksana

COL_KATEGORI_TES = 1
COL_NAMA_MODUL = 2
COL_NOMOR_SKENARIO = 3
COL_NOMOR_KASUS_TES = 4
COL_LANGKAH_TES = 5
COL_HASIL_DIHARAPKAN = 6
COL_HASIL_AKTUAL = 7
COL_REMARKS = 8
COL_TANGGAL = 9
COL_JENIS_SCRIPT = 10
COL_PELAKSANA = 11

# Definisi 6 section API di Lampiran 7C
# Format: (nama_lampiran, jumlah_skenario)
LAMPIRAN_SECTIONS = [
    ("API Balance Inquiry", 11),
    ("Intrabank Transfer", 12),
    ("Interbank Transfer", 13),
    ("API RTGS Transfer", 13),
    ("API SKNBI Transfer", 13),
    ("API Virtual Account", 33),
]

# Mapping dari section header di UAT Script ke section di Lampiran 7C
# Format: (header_uat, skenario_prefix, target_lampiran_section, is_fill_empty_only)
UAT_TO_LAMPIRAN_MAPPING = [
    ("Balance Services", "1", "API Balance Inquiry", False),
    ("Intrabank Transfer", "2", "Intrabank Transfer", False),
    ("Interbank Transfer", "3", "Interbank Transfer", False),
    # Catatan: "Interbank Transfer via BI FAST" (skenario 4.x) SENGAJA TIDAK
    # dipindahkan ke Lampiran 7C sesuai keputusan bisnis.
    ("RTGS Transfer", "5", "API RTGS Transfer", False),
    ("SKNBI Transfer", "6", "API SKNBI Transfer", False),
    ("Transfer VA", "7", "API Virtual Account", False),
    # Catatan: "Transfer VA Prima" (skenario 8.x) dan "Transfer VA BI FAST"
    # (skenario 9.x) SENGAJA TIDAK dipindahkan ke Lampiran 7C sesuai keputusan
    # bisnis (sama seperti Interbank Transfer via BI FAST).
]




# ============================================================
# EXCEL READER
# ============================================================

def get_cell_value(row, index):
    """Safely get cell value from a row tuple by index."""
    if index < len(row) and row[index] is not None:
        return str(row[index]).strip()
    return ""


def detect_section_header(row):
    """
    Detect if a row is a section header.

    Section header rules:
    - Kolom E (Nomor Kasus Tes, index 4) HARUS KOSONG
    - Kolom F (Langkah Tes, index 5) HARUS KOSONG
    - Teks section muncul di kolom B (index 1) atau kolom C (index 2)

    Returns:
        str: Detected section name, or None if not a section header
    """
    # Kolom E (Nomor Kasus Tes) harus kosong
    nomor_kasus = get_cell_value(row, COL_NOMOR_KASUS_TES)
    if nomor_kasus:
        return None

    # Kolom F (Langkah Tes) harus kosong
    langkah_tes = get_cell_value(row, COL_LANGKAH_TES)
    if langkah_tes:
        return None

    # Cek teks di kolom B atau C (gabung untuk pemeriksaan toleran)
    kategori_tes = get_cell_value(row, COL_KATEGORI_TES)
    nama_modul = get_cell_value(row, COL_NAMA_MODUL)
    t = (kategori_tes + " " + nama_modul).lower()

    # Deteksi TOLERAN berbasis sub-kata kunci (tidak terpaku string persis),
    # supaya variasi penulisan antar mitra tetap terdeteksi. Contoh yang kini
    # dikenali untuk Virtual Account: "Transfer VA", "Virtual Account",
    # "Transfer Virtual Account", "VA Transfer", dsb.
    # Urutan pengecekan: dari yang PALING SPESIFIK ke umum.

    has_bifast = ("bi fast" in t) or ("bifast" in t) or ("bi-fast" in t)
    has_prima = "prima" in t
    # "va" sebagai kata utuh, atau frasa "virtual account"
    has_va = ("virtual account" in t) or bool(re.search(r"\bva\b", t))

    # 1) Virtual Account - varian (paling spesifik dulu)
    if has_va and has_prima:
        return "Transfer VA Prima"
    if has_va and has_bifast:
        return "Transfer VA BI FAST"
    if has_va:
        return "Transfer VA"

    # 2) Interbank - varian
    if "interbank" in t and has_bifast:
        return "Interbank Transfer via BI FAST"
    if "interbank" in t:
        return "Interbank Transfer"

    # 3) Section lain (kata kunci inti)
    if "balance" in t:
        return "Balance Services"
    if "intrabank" in t:
        return "Intrabank Transfer"
    if "rtgs" in t:
        return "RTGS Transfer"
    if "sknbi" in t or re.search(r"\bskn\b", t):
        return "SKNBI Transfer"

    return None


def read_uat_script(filepath):
    """
    Membaca file UAT Script.xlsx dan mengembalikan data terstruktur.

    Kolom di UAT Script dimulai dari kolom B (kolom A kosong):
    - B (1): Kategori Tes
    - C (2): Nama Modul
    - D (3): Nomor Skenario
    - E (4): Nomor Kasus Tes
    - F (5): Langkah Tes
    - G (6): Hasil yang diharapkan
    - H (7): Hasil Aktual
    - I (8): Remarks
    - J (9): Tanggal Pelaksanaan
    - K (10): Jenis Script
    - L (11): Pelaksana

    Returns:
        dict: {
            section_header: [
                {
                    'nama_modul': str,
                    'nomor_skenario': str,
                    'nomor_kasus_tes': str,
                    'langkah_tes': str,
                    'hasil_diharapkan': str,
                    'hasil_aktual': str,
                    'remarks': str,
                }
            ]
        }
    """
    wb = load_workbook(filepath, read_only=True, data_only=True)

    if UAT_SHEET_NAME not in wb.sheetnames:
        # Coba cari sheet dengan nama yang mirip
        for sheet_name in wb.sheetnames:
            if "uat" in sheet_name.lower() or "script" in sheet_name.lower():
                ws = wb[sheet_name]
                break
        else:
            ws = wb.active
            print(f"  [WARNING] Sheet '{UAT_SHEET_NAME}' tidak ditemukan. Menggunakan sheet: {ws.title}")
    else:
        ws = wb[UAT_SHEET_NAME]

    data = {}
    current_section = None
    rows_list = list(ws.iter_rows(min_row=1, values_only=True))

    # Identifikasi header row (baris dengan label kolom)
    # Cek kolom B (index 1) untuk label "Kategori Tes" atau similar
    header_row_idx = 0
    for idx, row in enumerate(rows_list):
        if not row or len(row) <= COL_KATEGORI_TES:
            continue
        cell_val = get_cell_value(row, COL_KATEGORI_TES)
        if cell_val.lower() in ['kategori tes', 'kategori', 'no']:
            header_row_idx = idx
            break

    # Process data rows (after header)
    for idx in range(header_row_idx + 1, len(rows_list)):
        row = rows_list[idx]
        if not row or all(cell is None or str(cell).strip() == '' for cell in row):
            continue

        # Cek apakah ini section header
        detected_section = detect_section_header(row)
        if detected_section:
            current_section = detected_section
            if current_section not in data:
                data[current_section] = []
            continue

        # Jika belum ada section, skip
        if current_section is None:
            continue

        # Parse data row - harus punya Nomor Kasus Tes (kolom E, index 4)
        nomor_kasus = get_cell_value(row, COL_NOMOR_KASUS_TES)
        if not nomor_kasus:
            continue

        row_data = {
            'nama_modul': get_cell_value(row, COL_NAMA_MODUL),
            'nomor_skenario': get_cell_value(row, COL_NOMOR_SKENARIO),
            'nomor_kasus_tes': nomor_kasus,
            'langkah_tes': get_cell_value(row, COL_LANGKAH_TES),
            'hasil_diharapkan': get_cell_value(row, COL_HASIL_DIHARAPKAN),
            'hasil_aktual': get_cell_value(row, COL_HASIL_AKTUAL),
            'remarks': get_cell_value(row, COL_REMARKS),
        }

        data[current_section].append(row_data)

    wb.close()
    return data


# ============================================================
# REMARKS PARSER
# ============================================================

import json


# Parser Remarks dibuat FLEKSIBEL: hanya berpatokan pada KATA KUNCI inti
# (url / header / request / response), tidak terpaku pada kata pengiring.
# Jadi semua variasi mitra dikenali, contoh:
#   URL:  |  URL Endpoint:  |  Request URL:  |  URL Request:
#   Headers:  |  Header:  |  Header Request:  |  Request headers:
#   Request Body:  |  Request body:  |  Body:  |  Payload:
#   Response:  |  Response Body:  |  Response body:

# Batas panjang sebuah baris agar masih dianggap "label" (bukan isi data).
_LABEL_MAX_LEN = 40

# Penanda blok yang kadang MENEMPEL di tengah/akhir baris (mis. "}Response body:").
# Kita sisipkan newline sebelum penanda ini agar terbaca sebagai label terpisah.
_INLINE_MARKER_REGEX = re.compile(
    r"(?i)(?<=[}\]])"
    r"(?=(?:request|url|endpoint|header|headers|body|payload|response)"
    r"(?:[ \t]+\w+){0,2}[ \t]*:)"
)


def _normalize_inline_markers(text):
    """
    Sisipkan newline sebelum penanda blok yang menempel di akhir baris
    sebelumnya, contoh: '...}Response body: {...}' -> '...}\nResponse body: {...}'.
    Ini membuat parser per-baris dapat mengenalinya sebagai label.
    """
    return _INLINE_MARKER_REGEX.sub("\n", text)


def _classify_label(line):
    """
    Klasifikasikan sebuah baris label penanda ke salah satu blok:
    'url', 'headers', 'request_body', 'response'. Mengembalikan (nama, sisa_teks)
    di mana sisa_teks adalah teks setelah tanda ':' pada baris yang sama
    (biasanya kosong, tapi kadang URL menempel: "URL: https://...").

    Aturan (urut prioritas agar tidak salah klasifikasi):
      1. mengandung "response"           -> response
      2. mengandung "url" atau "endpoint"-> url
      3. mengandung "header"             -> headers
      4. mengandung "request"/"body"/"payload" -> request_body

    Baris hanya dianggap label jika:
      - mengandung tanda ':'
      - bagian SEBELUM ':' pendek (<= _LABEL_MAX_LEN) dan tidak berisi '{'/'['
        (supaya baris data JSON tidak salah dikira label).

    Jika bukan label, kembalikan (None, None).
    """
    if ":" not in line:
        return (None, None)

    before, after = line.split(":", 1)
    key = before.strip().lower()

    # Tolak kalau bagian sebelum ':' terlalu panjang atau tampak seperti data
    if len(key) == 0 or len(key) > _LABEL_MAX_LEN:
        return (None, None)
    if "{" in key or "[" in key or '"' in key:
        return (None, None)

    if "response" in key:
        return ("response", after.strip())
    if "url" in key or "endpoint" in key:
        return ("url", after.strip())
    if "header" in key:
        return ("headers", after.strip())
    if "request" in key or "body" in key or "payload" in key:
        return ("request_body", after.strip())

    return (None, None)


def _parse_remarks_blocks(text):
    """
    Pecah isi Remarks menjadi dict blok: {url, headers, request_body, response}.

    Bekerja per-baris: cari baris yang merupakan LABEL (via _classify_label),
    lalu kumpulkan semua baris berikutnya sebagai isi blok sampai ketemu label
    berikutnya. Pendekatan berbasis kata kunci ini tahan terhadap variasi
    penamaan antar mitra.
    """
    lines = text.split("\n")

    # Temukan indeks baris yang merupakan label + jenis bloknya
    label_positions = []  # list of (idx, name, inline_rest)
    for idx, line in enumerate(lines):
        name, rest = _classify_label(line)
        if name is not None:
            label_positions.append((idx, name, rest))

    blocks = {}
    for i, (idx, name, inline_rest) in enumerate(label_positions):
        start = idx + 1
        end = label_positions[i + 1][0] if i + 1 < len(label_positions) else len(lines)
        body_lines = lines[start:end]
        content = "\n".join(body_lines).strip()
        # Jika ada teks menempel di baris label (mis. "URL: https://..."),
        # gabungkan di depan.
        if inline_rest:
            content = (inline_rest + ("\n" + content if content else "")).strip()
        # Ambil kemunculan pertama tiap blok
        if name not in blocks and content:
            blocks[name] = content

    return blocks


def _format_headers(headers_raw):
    """
    Format blok Headers dari UAT Script ke format array Lampiran 7C.

    Mendukung 2 bentuk input:
      Bentuk A (list baris polos):
        [
        Content-Type: application/json
        Authorization: Bearer xxx
        ]
      Bentuk B (JSON object, nilai berupa array):
        {
          "Authorization": ["Bearer xxx"],
          "Content-Type": ["application/json"]
        }

    Output (format Lampiran 7C, beautify - satu header per baris):
        [
          "Content-Type=application/json",
          "Authorization=Bearer xxx"
        ]
    """
    pairs = []
    raw = headers_raw.strip()

    # Coba parse sebagai JSON dulu (bisa object ATAU array of strings)
    parsed = None
    try:
        parsed = json.loads(raw)
    except (ValueError, TypeError):
        parsed = None

    if isinstance(parsed, dict):
        # Bentuk B: {"Authorization": ["Bearer xxx"], ...}
        for key, val in parsed.items():
            if isinstance(val, list):
                val = ", ".join(str(v) for v in val)
            pairs.append(f'{key}={val}')
    elif isinstance(parsed, list):
        # Bentuk C: ["Content-Type=application/json", "Authorization=Bearer xxx"]
        # Item sudah "Key=Value" -> pakai apa adanya (jangan bungkus kutip lagi).
        for item in parsed:
            pairs.append(str(item))
    else:
        # Bentuk A: parsing per baris polos
        for line in raw.split("\n"):
            s = line.strip().rstrip(",")
            if not s or s in ("[", "]", "{", "}"):
                continue
            # Buang kutip pembungkus di sekeliling item bila ada
            # (mis. baris '"Content-Type=application/json"').
            if len(s) >= 2 and s[0] == '"' and s[-1] == '"':
                s = s[1:-1]
            if "=" in s and ":" not in s.split("=", 1)[0]:
                # Sudah format Key=Value
                pairs.append(s)
            elif ":" in s:
                key, val = s.split(":", 1)
                pairs.append(f'{key.strip().strip(chr(34))}={val.strip().strip(chr(34))}')
            else:
                pairs.append(s)

    if not pairs:
        return ""

    # Beautify: satu header per baris, indentasi 2 spasi, dipisah koma.
    #   [
    #     "Key=Value",
    #     "Key=Value"
    #   ]
    out = ["["]
    for i, p in enumerate(pairs):
        comma = "," if i < len(pairs) - 1 else ""
        out.append(f'  "{p}"{comma}')
    out.append("]")
    return "\n".join(out)


def _pretty_json(raw):
    """
    Pretty-print (beautify) JSON dengan indentasi 2 spasi jika bisa di-parse,
    agar mudah dibaca pada laporan. Jika gagal parse, kembalikan teks asli
    apa adanya (sudah di-strip).
    """
    raw = raw.strip()
    if not raw:
        return ""
    try:
        obj = json.loads(raw)
        return json.dumps(obj, indent=2, ensure_ascii=False)
    except (ValueError, TypeError):
        # Bukan JSON valid -> kembalikan apa adanya
        return raw


def _compress_json(raw):
    """
    Compress JSON menjadi satu baris tanpa spasi jika bisa di-parse.
    Jika gagal parse, rapikan jadi satu baris (buang newline & spasi ganda).
    Dipakai khusus untuk Response Body.
    """
    raw = raw.strip()
    if not raw:
        return ""
    try:
        obj = json.loads(raw)
        return json.dumps(obj, separators=(",", ":"), ensure_ascii=False)
    except (ValueError, TypeError):
        return re.sub(r"\s+", " ", raw).strip()


def split_request_response(remarks_text):
    """
    Memisahkan DAN memformat isi kolom Remarks menjadi bagian Request dan
    Response sesuai format Lampiran 7C.

    Format Remarks (sumber, dari UAT Script):
        URL:
        <url>

        Headers:
        [ Key: Value ... ]

        Request Body:
        { ... }

        Response:
        { ... }

    Format hasil (Lampiran 7C):
        Request:
            URL Endpoint:
            <url>

            Header Request:
            [
              "Key=Value",
              ...
            ]

            Request Body:
            { ...pretty... }

        Response:
            Response Body:
            { ... }

    Returns:
        tuple: (request_content, response_content) keduanya sudah di-strip.
    """
    if not remarks_text or remarks_text.strip().lower() == "none":
        return "", ""

    text = remarks_text.replace("\r\n", "\n").replace("\r", "\n")
    text = _normalize_inline_markers(text)

    blocks = _parse_remarks_blocks(text)

    url_raw = blocks.get("url", "")
    headers_raw = blocks.get("headers", "")
    body_raw = blocks.get("request_body", "")
    response_raw = blocks.get("response", "")

    # URL kadang diawali HTTP method (contoh: "POST https://..."). Ambil apa
    # adanya, hanya bersihkan whitespace.
    url_raw = url_raw.strip()

    # Susun bagian Request
    request_blocks = []
    if url_raw:
        request_blocks.append(f"URL Endpoint:\n{url_raw}")
    if headers_raw:
        request_blocks.append(f"Header Request:\n{_format_headers(headers_raw)}")
    if body_raw:
        request_blocks.append(f"Request Body:\n{_pretty_json(body_raw)}")

    request_part = "\n\n".join(request_blocks).strip()

    # Susun bagian Response (Response Body di-compress, satu baris)
    if response_raw:
        response_part = f"Response Body:\n{_compress_json(response_raw)}"
    else:
        response_part = ""

    # Fallback: jika format tidak dikenali sama sekali, pisahkan sederhana
    # pada penanda "Response" agar data tidak hilang.
    if not request_part and not response_part:
        m = re.search(r"(?i)response(?:\s*body)?\s*:", text)
        if m:
            request_part = text[:m.start()].strip()
            response_part = text[m.end():].strip()
        else:
            request_part = text.strip()

    return request_part, response_part


# ============================================================
# MAPPING LOGIC
# ============================================================

def extract_sub_number(nomor_kasus_tes):
    """
    Extract sub-number dari Nomor Kasus Tes.
    Contoh: "2.3" -> 3, "7.15" -> 15, "1.1" -> 1

    Args:
        nomor_kasus_tes: String nomor kasus (e.g., "2.3")

    Returns:
        int: Sub-number (1-based index untuk baris di Lampiran 7C)
    """
    parts = nomor_kasus_tes.split('.')
    if len(parts) >= 2:
        try:
            return int(parts[1])
        except ValueError:
            return None
    return None


def map_result_value(hasil_aktual):
    """
    Map Hasil Aktual ke Result value di Lampiran 7C.

    Mengacu pada legend "KET" di UAT Script:
    - "Berhasil"          -> "PASS"     (sesuai hasil yang diharapkan)
    - "Gagal"             -> "NOT PASS" (tidak sesuai hasil yang diharapkan)
    - "Tidak dites"       -> "N/A"      (di luar scope / tidak ada alat pendukung)
    - "Belum dites"       -> "N/A"      (script belum waktunya dites)
    - "Siap dites"        -> "N/A"      (menunggu update fixing dari vendor)
    - "Butuh Konfirmasi"  -> "N/A"      (butuh konfirmasi user/vendor)
    - Lainnya             -> nilai asli
    """
    if not hasil_aktual:
        return ""
    lower = hasil_aktual.lower().strip()
    if lower == "berhasil":
        return "PASS"
    elif lower == "gagal":
        return "NOT PASS"
    elif lower in ("tidak dites", "belum dites", "siap dites", "butuh konfirmasi"):
        return "N/A"
    return hasil_aktual


def _is_valid_json(raw):
    """True jika teks bisa di-parse sebagai JSON."""
    raw = (raw or "").strip()
    if not raw:
        return False
    try:
        json.loads(raw)
        return True
    except (ValueError, TypeError):
        return False


def detect_anomalies(row_data):
    """
    Deteksi kondisi ABNORMAL pada satu baris UAT agar bisa ditampilkan sebagai
    peringatan (tidak mengubah data apa pun - hanya memberi tahu).

    PENTING: fungsi ini TIDAK mengubah data apa pun. Hanya melaporkan lokasi
    potensi masalah agar bisa di-cross-check manual oleh pengguna.

    Kondisi yang dideteksi (khusus baris dengan Hasil Aktual = "Berhasil"):
      1. Remarks kosong padahal hasil Berhasil
      2. Ada bagian log yang hilang (URL / Header / Request Body / Response)
      3. Request Body atau Response Body bukan JSON valid (syntax rusak)
      4. Response Body kosong padahal hasil Berhasil
      5. Ada item Header yang tidak berbentuk "Key=Value" (tidak wajar)

    Args:
        row_data: dict baris UAT (punya 'nomor_kasus_tes', 'hasil_aktual', 'remarks')

    Returns:
        list[str]: daftar pesan peringatan (kosong jika tidak ada anomali)
    """
    warnings = []
    kasus = row_data.get('nomor_kasus_tes', '?')
    hasil = str(row_data.get('hasil_aktual', '')).strip().lower()
    remarks = row_data.get('remarks', '') or ""

    # Hanya periksa baris yang seharusnya punya log (Berhasil)
    if hasil != "berhasil":
        return warnings

    if not remarks.strip() or remarks.strip().lower() == "none":
        warnings.append(f"Kasus {kasus}: hasil 'Berhasil' tetapi kolom Remarks kosong.")
        return warnings

    text = remarks.replace("\r\n", "\n").replace("\r", "\n")
    text = _normalize_inline_markers(text)
    blocks = _parse_remarks_blocks(text)

    # (2) Bagian yang hilang
    if not blocks.get("url", "").strip():
        warnings.append(f"Kasus {kasus}: bagian URL tidak ditemukan pada Remarks.")
    if not blocks.get("headers", "").strip():
        warnings.append(f"Kasus {kasus}: bagian Header tidak ditemukan pada Remarks.")
    if not blocks.get("request_body", "").strip():
        warnings.append(f"Kasus {kasus}: bagian Request Body tidak ditemukan pada Remarks.")
    if not blocks.get("response", "").strip():
        warnings.append(f"Kasus {kasus}: bagian Response tidak ditemukan pada Remarks.")

    # (3) JSON tidak valid (hanya cek jika bagiannya ada)
    body_raw = blocks.get("request_body", "")
    if body_raw.strip() and not _is_valid_json(body_raw):
        warnings.append(
            f"Kasus {kasus}: Request Body bukan JSON valid (format/syntax) - mohon cek manual."
        )

    resp_raw = blocks.get("response", "")
    if resp_raw.strip():
        if not _is_valid_json(resp_raw):
            warnings.append(
                f"Kasus {kasus}: Response Body bukan JSON valid (format/syntax) - mohon cek manual."
            )
    else:
        # (4) Response kosong padahal Berhasil
        warnings.append(
            f"Kasus {kasus}: hasil 'Berhasil' tetapi Response Body kosong - mohon cek manual."
        )

    # (5) Header ada item yang tidak berbentuk Key=Value
    headers_raw = blocks.get("headers", "")
    if headers_raw.strip():
        for item in _iter_header_items(headers_raw):
            # Item wajar bila mengandung pemisah ':' atau '=' antara key & value
            if ("=" not in item) and (":" not in item):
                preview = item[:40] + ("..." if len(item) > 40 else "")
                warnings.append(
                    f"Kasus {kasus}: ada Header tidak wajar (bukan Key=Value): '{preview}' - mohon cek manual."
                )
                break  # cukup satu peringatan header per kasus

    return warnings


def _iter_header_items(headers_raw):
    """
    Kembalikan daftar item header (string) dari blok Header mentah, mendukung
    JSON object, JSON array, maupun baris polos. Hanya untuk PEMERIKSAAN
    (tidak mengubah data).
    """
    raw = headers_raw.strip()
    items = []
    parsed = None
    try:
        parsed = json.loads(raw)
    except (ValueError, TypeError):
        parsed = None

    if isinstance(parsed, dict):
        for k, v in parsed.items():
            if isinstance(v, list):
                v = ", ".join(str(x) for x in v)
            items.append(f"{k}: {v}")
    elif isinstance(parsed, list):
        items = [str(x) for x in parsed]
    else:
        for line in raw.split("\n"):
            s = line.strip().rstrip(",")
            if not s or s in ("[", "]", "{", "}"):
                continue
            if len(s) >= 2 and s[0] == '"' and s[-1] == '"':
                s = s[1:-1]
            items.append(s)
    return items


def map_uat_to_lampiran(uat_data, collect_warnings=False):
    """
    Map data dari UAT Script ke struktur Lampiran 7C.

    Mengimplementasikan logika "fill empty only" untuk section yang di-share.

    Args:
        uat_data: Dict hasil dari read_uat_script()
        collect_warnings: jika True, kembalikan tuple (lampiran_data, warnings)

    Args:
        uat_data: Dict hasil dari read_uat_script()

    Returns:
        dict: {
            section_name: [
                {
                    'no': int,
                    'service': str,
                    'scenario': str,
                    'expected_result': str,
                    'request': str,
                    'response': str,
                    'result': str,
                    'notes': str,
                }
                atau None (jika row belum terisi)
            ]
        }
    """
    # Inisialisasi struktur Lampiran 7C dengan None untuk setiap row
    lampiran_data = {}
    for section_name, count in LAMPIRAN_SECTIONS:
        lampiran_data[section_name] = [None] * count

    warnings = []

    # Proses mapping berdasarkan urutan prioritas
    for uat_section, _, target_section, fill_empty_only in UAT_TO_LAMPIRAN_MAPPING:
        if uat_section not in uat_data:
            continue

        rows = uat_data[uat_section]
        target_count = None
        for sec_name, sec_count in LAMPIRAN_SECTIONS:
            if sec_name == target_section:
                target_count = sec_count
                break

        if target_count is None:
            continue

        for row_data in rows:
            sub_num = extract_sub_number(row_data['nomor_kasus_tes'])
            if sub_num is None or sub_num < 1 or sub_num > target_count:
                continue

            row_idx = sub_num - 1  # Convert ke 0-based index

            # Cek apakah harus fill empty only
            if fill_empty_only and lampiran_data[target_section][row_idx] is not None:
                continue

            # Parse remarks
            hasil_aktual = row_data['hasil_aktual']
            remarks_text = row_data['remarks']

            request_content = ""
            response_content = ""
            notes = ""
            result_value = map_result_value(hasil_aktual)

            # Service diambil dari kolom Nama Modul
            service = row_data.get('nama_modul', '') or target_section

            # Cek kondisi khusus
            if hasil_aktual.lower().strip() in ("tidak dites", "belum dites"):
                # Tidak dites / Belum dites: baris tetap ditampilkan (Result N/A).
                # Kolom Request & Response dikosongkan, dan isi Remark dipindahkan
                # ke kolom Notes (apa adanya).
                request_content = ""
                response_content = ""
                if remarks_text and remarks_text.strip().lower() != "none":
                    notes = remarks_text.strip()
                else:
                    notes = ""
            elif not remarks_text or remarks_text.lower() == "none":
                # Remarks kosong: Request dan Response kosong
                notes = ""
            else:
                # Pisahkan Remarks menjadi Request (URL+Headers+Request Body)
                # dan Response (isi setelah penanda "Response:")
                request_content, response_content = split_request_response(remarks_text)

            # Kumpulkan peringatan anomali (tidak mengubah data)
            warnings.extend(detect_anomalies(row_data))

            lampiran_data[target_section][row_idx] = {
                'no': sub_num,
                'no_full': row_data['nomor_kasus_tes'],
                'service': service,
                'scenario': row_data.get('langkah_tes', '') or f"Skenario {row_data['nomor_kasus_tes']}",
                'expected_result': row_data.get('hasil_diharapkan', ''),
                'request': request_content,
                'response': response_content,
                'result': result_value,
                'notes': notes,
            }

    if collect_warnings:
        return lampiran_data, warnings
    return lampiran_data


# ============================================================
# WORD DOCUMENT GENERATOR
# ============================================================

def _set_page(section):
    """
    Set halaman LANDSCAPE dengan ukuran & margin sesuai contoh Lampiran 7C
    (Letter landscape: 27.94 x 21.59 cm).
    """
    from docx.enum.section import WD_ORIENT

    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(27.94)
    section.page_height = Cm(21.59)
    # Margin sesuai contoh
    section.left_margin = Cm(0.9)
    section.right_margin = Cm(1.27)
    section.top_margin = Cm(1.27)
    section.bottom_margin = Cm(1.27)


def _set_table_fixed_layout(table, col_widths_twips):
    """
    Kunci lebar kolom tabel agar Word TIDAK meng-autofit (yang menyebabkan
    teks pecah vertikal). Caranya:
      1. Set tblLayout = fixed
      2. Set lebar total tabel (tblW)
      3. Definisikan <w:tblGrid> berisi lebar tiap kolom (twips)

    Args:
        table: objek tabel docx
        col_widths_twips: list lebar kolom dalam twips
    """
    tbl = table._tbl
    tblPr = tbl.tblPr

    # 1. Layout fixed (bukan autofit)
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    # 2. Lebar total tabel
    total = sum(col_widths_twips)
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")

    # 3. Definisikan grid kolom (buang yang lama jika ada)
    old_grid = tbl.find(qn("w:tblGrid"))
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for w in col_widths_twips:
        gc = OxmlElement("w:gridCol")
        gc.set(qn("w:w"), str(w))
        grid.append(gc)
    # tblGrid harus diletakkan tepat setelah tblPr
    tblPr.addnext(grid)


def _set_cell_background(cell, hex_color):
    """
    Memberi warna latar (shading) pada sebuah sel tabel.

    python-docx tidak menyediakan API langsung untuk shading sel, jadi kita
    menyisipkan elemen <w:shd> ke properti sel (tcPr) secara manual.

    Args:
        cell: objek sel tabel (docx table cell)
        hex_color: warna hex tanpa '#', contoh "ED7D31"
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _write_cell_rich(cell, text):
    """
    Menulis teks ke sel, di mana LABEL tertentu (URL Endpoint:, Header Request:,
    Request Body:, Response Body:) dibuat BOLD sedangkan isinya normal.

    Teks multi-baris: setiap baris jadi satu paragraf. Jika suatu baris diawali
    salah satu label, bagian label dicetak tebal dan sisanya (jika ada di baris
    yang sama) normal.
    """
    # Kosongkan paragraf default sel
    cell.text = ""
    lines = text.split("\n")

    first = True
    for line in lines:
        p = cell.paragraphs[0] if first else cell.add_paragraph()
        first = False

        matched_label = None
        for label in BOLD_LABELS:
            if line.strip() == label or line.startswith(label):
                matched_label = label
                break

        if matched_label:
            run_label = p.add_run(matched_label)
            run_label.font.bold = True
            run_label.font.name = FONT_NAME
            sisa = line[len(matched_label):]
            if sisa:
                run_rest = p.add_run(sisa)
                run_rest.font.name = FONT_NAME
        else:
            run = p.add_run(line)
            run.font.name = FONT_NAME


def build_lampiran_document(lampiran_data):
    """
    Membangun dokumen Word Lampiran 7C dari data hasil mapping dan
    mengembalikan objek Document (tanpa menyimpan ke disk).

    Fungsi ini dipakai bersama oleh CLI (main.py) maupun aplikasi web (app.py).

    Args:
        lampiran_data: Dict hasil dari map_uat_to_lampiran()

    Returns:
        docx.Document: dokumen yang siap disimpan.
    """
    doc = Document()

    # Halaman landscape dengan margin sesuai contoh
    _set_page(doc.sections[0])

    # Set default font seluruh dokumen (Calibri)
    style = doc.styles['Normal']
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(9)
    # Pastikan berlaku untuk semua jenis skrip (latin, cs, dll.)
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), FONT_NAME)
    rfonts.set(qn("w:hAnsi"), FONT_NAME)
    rfonts.set(qn("w:cs"), FONT_NAME)

    # Hanya render section yang benar-benar punya data hasil tes.
    # Section yang seluruh barisnya kosong (None) tidak ditampilkan.
    sections_to_render = []
    for section_name, section_count in LAMPIRAN_SECTIONS:
        section_data = lampiran_data.get(section_name, [])
        filled_rows = [r for r in section_data if r is not None]
        if filled_rows:
            sections_to_render.append((section_name, filled_rows))

    if not sections_to_render:
        # Tidak ada satupun layanan yang dites
        doc.add_paragraph("Tidak ada data hasil UAT yang ditemukan pada UAT Script.")
        return doc

    # Judul dokumen (format resmi Lampiran 7.C) - HANYA DITAMPILKAN SEKALI
    # di halaman 1 (sebelum section pertama), tidak diulang tiap section.
    title1 = doc.add_paragraph()
    title1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = title1.add_run("Lampiran 7.C")
    run1.font.name = FONT_NAME
    run1.font.size = Pt(12)
    run1.font.bold = True

    title2 = doc.add_paragraph()
    title2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = title2.add_run("Skenario dan Hasil Uji Fungsionalitas")
    run2.font.name = FONT_NAME
    run2.font.size = Pt(11)
    run2.font.bold = True

    doc.add_paragraph()

    # Buat setiap section yang ada datanya
    for section_idx, (section_name, filled_rows) in enumerate(sections_to_render):
        # Tambah page break sebelum section (kecuali section pertama)
        if section_idx > 0:
            doc.add_page_break()

        # Info penyedia & layanan - paragraf rapat (single line, spacing 0pt)
        info_lines = [
            f"Nama Penyedia Layanan : {NAMA_PENYEDIA_LAYANAN}",
            f"Nama Pengguna Layanan : {NAMA_PENGGUNA_LAYANAN}",
            f"Nama Layanan API      : {section_name}",
            f"Tanggal Pengujian     : {TANGGAL_PENGUJIAN}",
        ]
        for line in info_lines:
            p = doc.add_paragraph(line)
            pf = p.paragraph_format
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.line_spacing = 1.5

        doc.add_paragraph()

        # Buat tabel
        # Kolom: No, Service, Scenario, Expected Result, Request, Response, Result, Notes
        table = doc.add_table(rows=1, cols=8)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        # Matikan auto-fit (Word: "Fixed column width")
        table.autofit = False
        # Kunci lebar kolom (fixed layout) agar teks tidak pecah vertikal
        _set_table_fixed_layout(table, COLUMN_WIDTHS_TWIPS)

        # Header row - latar oranye, teks bold, rata tengah
        header_cells = table.rows[0].cells
        headers = ['No', 'Service', 'Scenario', 'Expected Result',
                   'Request', 'Response', 'Result', 'Notes']
        for i, header_text in enumerate(headers):
            cell = header_cells[i]
            cell.text = header_text
            _set_cell_background(cell, HEADER_FILL_COLOR)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.name = FONT_NAME

        # Data rows - hanya baris yang benar-benar dites (skip baris kosong)
        # Kolom No menampilkan nomor kasus LENGKAP, mis. "1.1", "2.12", "7.33".
        for row_data in filled_rows:
            row_cells = table.add_row().cells
            values = [
                str(row_data.get('no_full', row_data['no'])),
                row_data['service'],
                row_data['scenario'],
                row_data['expected_result'],
                row_data['request'],
                row_data['response'],
                row_data['result'],
                row_data['notes'],
            ]
            # Baris Tidak dites / Belum dites (Result = N/A) diblok kuning
            is_na = str(row_data.get('result', '')).strip().upper() == "N/A"

            for ci, val in enumerate(values):
                cell = row_cells[ci]
                # Kolom Request (4) & Response (5): label dibuat bold, isi normal
                if ci in (4, 5):
                    _write_cell_rich(cell, val)
                else:
                    cell.text = val
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                if is_na:
                    _set_cell_background(cell, NA_FILL_COLOR)
                # Kolom tertentu rata tengah, sisanya rata kiri
                align = (WD_ALIGN_PARAGRAPH.CENTER
                         if ci in CENTER_ALIGNED_COLUMNS
                         else WD_ALIGN_PARAGRAPH.LEFT)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = align

        # Set lebar kolom di tiap sel (memperkuat fixed layout dari tblGrid).
        # docx.Cm butuh EMU; twips -> cm: twips / 567
        for row in table.rows:
            for ci, width_twips in enumerate(COLUMN_WIDTHS_TWIPS):
                row.cells[ci].width = Cm(width_twips / 567.0)

    return doc


def create_lampiran_document(lampiran_data, output_path):
    """
    Membuat dokumen Word Lampiran 7C dan menyimpannya ke output_path.
    (Wrapper tipis di atas build_lampiran_document untuk pemakaian CLI.)
    """
    doc = build_lampiran_document(lampiran_data)
    doc.save(str(output_path))
    print(f"  [OK] Dokumen berhasil disimpan: {output_path}")


# ============================================================
# API REUSABLE (dipakai aplikasi web)
# ============================================================

def convert_uat_to_lampiran(source):
    """
    Konversi UAT Script (Excel) menjadi dokumen Lampiran 7C.

    Fungsi tingkat tinggi yang menyatukan 3 langkah: baca -> mapping -> bangun
    dokumen. Cocok dipanggil dari aplikasi web.

    Args:
        source: path file (str/Path) ATAU objek file-like/bytes berisi .xlsx

    Returns:
        tuple: (doc, stats, warnings)
            doc      : docx.Document hasil konversi (belum disimpan)
            stats    : dict {section_name: jumlah_baris_terisi} untuk ringkasan
            warnings : list[str] daftar peringatan anomali (bisa kosong)
    """
    import io

    # openpyxl menerima path maupun file-like object. Jika berupa bytes,
    # bungkus dengan BytesIO.
    if isinstance(source, (bytes, bytearray)):
        source = io.BytesIO(source)

    uat_data = read_uat_script(source)
    lampiran_data, warnings = map_uat_to_lampiran(uat_data, collect_warnings=True)

    stats = {}
    for section_name, rows in lampiran_data.items():
        stats[section_name] = sum(1 for r in rows if r is not None)

    doc = build_lampiran_document(lampiran_data)
    return doc, stats, warnings


def convert_uat_to_lampiran_bytes(source):
    """
    Sama seperti convert_uat_to_lampiran(), tetapi mengembalikan dokumen dalam
    bentuk bytes (siap dikirim sebagai unduhan di aplikasi web).

    Returns:
        tuple: (docx_bytes, stats, warnings)
    """
    import io

    doc, stats, warnings = convert_uat_to_lampiran(source)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue(), stats, warnings


# ============================================================
# MAIN ENTRY POINT
# ============================================================

def main():
    """Fungsi utama untuk menjalankan automation Lampiran 7C."""
    print("=" * 60)
    print("  AUTOMATION LAMPIRAN 7C ASPI")
    print("  Konversi UAT Script (Excel) -> Lampiran 7C (Word)")
    print("=" * 60)
    print()

    # Cek file input
    uat_script_path = INPUT_DIR / UAT_SCRIPT_FILENAME
    if not uat_script_path.exists():
        print(f"  [ERROR] File UAT Script tidak ditemukan!")
        print(f"  Lokasi yang dicari: {uat_script_path.resolve()}")
        print()
        print("  Langkah yang harus dilakukan:")
        print(f"  1. Copy file UAT Script Anda ke folder: {INPUT_DIR.resolve()}")
        print(f"  2. Pastikan nama file: {UAT_SCRIPT_FILENAME}")
        print("  3. Jalankan ulang script ini")
        print()
        sys.exit(1)

    # Pastikan folder output ada
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Baca UAT Script
    print(f"  [1/3] Membaca file UAT Script: {uat_script_path}")
    try:
        uat_data = read_uat_script(uat_script_path)
    except Exception as e:
        print(f"  [ERROR] Gagal membaca file UAT Script: {e}")
        sys.exit(1)

    print(f"        Ditemukan {len(uat_data)} section:")
    for section_name, rows in uat_data.items():
        print(f"        - {section_name}: {len(rows)} baris data")
    print()

    # Mapping data
    print("  [2/3] Memproses mapping UAT Script -> Lampiran 7C...")
    lampiran_data = map_uat_to_lampiran(uat_data)

    # Hitung statistik
    total_filled = 0
    for section_name, rows in lampiran_data.items():
        filled = sum(1 for r in rows if r is not None)
        total_filled += filled
        section_count = next(c for n, c in LAMPIRAN_SECTIONS if n == section_name)
        print(f"        - {section_name}: {filled}/{section_count} baris terisi")
    print()

    # Generate dokumen Word
    output_path = OUTPUT_DIR / OUTPUT_FILENAME
    print(f"  [3/3] Membuat dokumen Lampiran 7C: {output_path}")
    try:
        create_lampiran_document(lampiran_data, output_path)
    except Exception as e:
        print(f"  [ERROR] Gagal membuat dokumen Word: {e}")
        sys.exit(1)

    print()
    print("  " + "=" * 56)
    print(f"  SELESAI! Total {total_filled} baris data berhasil dipindahkan.")
    print(f"  File output: {output_path.resolve()}")
    print("  " + "=" * 56)


if __name__ == "__main__":
    main()
